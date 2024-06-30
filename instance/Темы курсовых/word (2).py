from docx import Document

# Открываем исходный документ
doc_path = 'C:\D\учеба\Flask\stud_sait\instance\Темы курсовых работ на 2023.docx'
doc = Document(doc_path)

# Создаем новый документ
new_doc = Document()

# Функция для преобразования римских цифр в арабские
def roman_to_arabic(roman):
    roman_dict = {'I': 1, 'II': 2, 'III': 3, 'IV': 4}
    return str(roman_dict.get(roman, roman))

# Парсим документ и изменяем текст перед таблицами
i = 0
while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    text = para.text.strip()
    
    if text.startswith('Дисциплина'):
        # Извлекаем информацию о дисциплине, преподавателе и студентах
        discipline = text.split('«')[1].split('»')[0]
        teacher_info = text.split('(')[1].split(')')[0]
        teacher_name = ' '.join(part for part in teacher_info.split() if part[0].isupper())
        
        # Переходим к следующему параграфу для получения информации о курсе
        i += 1
        next_text = doc.paragraphs[i].text.strip()
        if 'студенты' in next_text:
            course_info = next_text.split('студенты')[1].split('курса')[0].strip()
            course_number = roman_to_arabic(course_info.split()[0])
            
            # Добавляем измененный текст в новый документ
            new_doc.add_paragraph(f'Дисциплина {discipline}')
            new_doc.add_paragraph(f'Преподаватель {teacher_name}')
            new_doc.add_paragraph(f'студенты {course_number}')
            print(f'Дисциплина {discipline}')
            print(f'Преподаватель {teacher_name}')
            print(f'студенты {course_number}')
            new_doc.add_paragraph(next_text)
        else:
            # Если следующая строка не содержит информации о студентах, возвращаемся обратно
            i -= 1

    new_doc.add_paragraph(text)
    i += 1

# Сохраняем новый документ
new_doc_path = 'C:\D\учеба\Flask\stud_sait\instance\Измененные_темы_курсовых_работ_3.docx'
new_doc.save(new_doc_path)

new_doc_path
