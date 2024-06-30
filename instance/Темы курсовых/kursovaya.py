import os
from docx import Document
import sqlite3
import re


def save_to_database(db_file, year, current_subject, current_teacher, current_course, table_data):
    conn = sqlite3.connect(db_file)
    c = conn.cursor()
    for student, coursework in table_data:
         c.execute("INSERT INTO List_coursework (year, subject, name_teacher, course, name_student, name_coursework) VALUES (?, ?, ?, ?, ?, ?)",
              (year, current_subject, current_teacher, current_course, student, coursework))

    conn.commit()
    conn.close()

def extract_year_from_filename(filename):
    basename = os.path.basename(filename)
    year = re.search(r'\b\d{4}\b', basename)
    if year:
        return year.group()
    else:
        raise ValueError("Year not found in filename: {}".format(filename))

def parse_word_file(docx_file, db_file):
    print("Parsing file:", docx_file)
    doc = Document(docx_file)
    year = extract_year_from_filename(docx_file)

    current_subject = ""
    current_teacher = ""
    current_course = ""

    table_started = False
    table_data = []
    table_index = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        print("Current text:", text)
        if text.startswith("Дисциплина"):
            current_subject = text.split("Дисциплина", 1)[1].strip()
            print("Subject:", current_subject)
        elif text.startswith("Преподаватель"):
            current_teacher = text.split("Преподаватель", 1)[1].strip()
            print("Teacher:", current_teacher)
        elif text.startswith("студенты"):
            current_course = text.split("студенты", 1)[1].strip()
            print("Course:", current_course)
            table_started = True
        elif table_started and doc.tables:
            table_data = parse_table(doc.tables[table_index])  # Парсинг таблицы
            if table_data:
                print("Processing table data...")
                print("Table data:", table_data)
                print("Saving to database...")
                save_to_database(db_file, year, current_subject, current_teacher, current_course, table_data)
            table_index += 1
            table_started = False



def parse_table(table):
    print("Parsing table...")
    if table:
        data = []
        for row in table.rows[1:]:  # Skip the header row
            cells = row.cells
            if len(cells) >= 3:
                name_student = cells[1].text.strip()
                name_coursework = cells[2].text.strip()
                data.append((name_student, name_coursework))
        print("Parsed data from table:", data)
        return data
    return None



# Получаем список всех .docx файлов в текущей директории
docx_files = [file for file in os.listdir() if file.endswith(".docx")]

for docx_file in docx_files:
    parse_word_file(docx_file, "sait.db")
