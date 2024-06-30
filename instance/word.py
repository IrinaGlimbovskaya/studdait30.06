"""import sqlite3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Функция для преобразования арабских цифр в римские
def arabic_to_roman(num):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
        ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
        ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num


# Подключение к базе данных SQLite
conn = sqlite3.connect('sait.db')
cursor = conn.cursor()

# Получение имени студента от пользователя
student_name_input = input("Введите имя студента: ")

# Извлечение данных по имени студента из базы данных
cursor.execute("SELECT course, faculty, name FROM studweb WHERE name = ?", (student_name_input,))
rows = cursor.fetchall()

if not rows:
    print("Студент с таким именем не найден.")
else:
    # Открытие существующего шаблонного документа
    template_doc_path = 'template.docx'
    doc = Document(template_doc_path)

    # Создание нового документа
    new_doc = Document()

    # Обработка и вставка данных в документ Word
for row in rows:
    course_number, group_number, student_name = row
    course_number_roman = arabic_to_roman(int(course_number))

    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.style = paragraph.style
        new_paragraph.alignment = paragraph.alignment

        for run in paragraph.runs:
            new_run = new_paragraph.add_run()
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name

            # Замена текста в run
            text = run.text
            text = text.replace("{COURSE_NUMBER}", course_number_roman)
            text = text.replace("{GROUP_NUMBER}", group_number)
            text = text.replace("{STUDENT_NAME}", student_name)

            # Добавление интервала перед и после текста
            text = f'\n{text.strip()}\n'

            new_run.text = text

    # Применение одинарного интервала к всему документу
    for paragraph in new_doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    # Сохранение нового документа с именем студента
    new_doc_path = f'Характеристика_{student_name_input}.docx'
    new_doc.save(new_doc_path)


    print(f"Документ сохранен по пути: {new_doc_path}")

# Закрытие соединения с базой данных
conn.close()
"""

import os
from docx import Document
import sqlite3

# Функция для преобразования арабских цифр в римские
def arabic_to_roman(num):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
        ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
        ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

# Подключение к базе данных SQLite
conn = sqlite3.connect('sait.db')
cursor = conn.cursor()

# Получение имени студента от пользователя
student_name_input = input("Введите имя студента: ")

# Получение данных из базы данных
cursor.execute("SELECT course, name FROM studweb WHERE name = ?", (student_name_input,))
rows = cursor.fetchall()

# Открытие шаблонного документа
template_doc = Document("template.docx")

# Вставка данных из базы данных в шаблонный документ
for row in rows:
    # Распаковка данных из строки
    course_number, student_name = row

    # Получение номера группы из другой таблицы
    cursor.execute("SELECT name FROM courses WHERE course = ?", (course_number,))
    group_number = cursor.fetchone()[0]
    print(group_number)
    # Подстановка данных в шаблонный документ
    for paragraph in template_doc.paragraphs:
        for run in paragraph.runs:
            if "{COURSE_NUMBER}" in run.text:
                run.text = run.text.replace("{COURSE_NUMBER}", arabic_to_roman(course_number))
            elif "{РФ}" in run.text:
                print(group_number)
                run.text = run.text.replace("{РФ}", group_number)
            elif "{STUDENT_NAME}" in run.text:
                run.text = run.text.replace("{STUDENT_NAME}", student_name)

# Сохранение нового документа
new_doc_path = f'Характеристика_{student_name_input}.docx'
template_doc.save(new_doc_path)
print(f"Документ сохранен по пути: {new_doc_path}")

# Закрытие соединения с базой данных
conn.close()
